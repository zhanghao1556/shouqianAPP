from __future__ import annotations

import json
import os
import re
import shutil
import subprocess
import sys
import textwrap
import zipfile
from datetime import date
from pathlib import Path
from typing import Iterable, Sequence

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "software-copyright"
SCREENSHOT_DIR = OUT / "screenshots"
QA_DIR = OUT / "qa"

SOFTWARE_NAME = "智能音频售前工程方案设计软件"
SHORT_NAME = "智能音频售前工具"
VERSION = "V2.0"
FULL_NAME = f"{SOFTWARE_NAME} {VERSION}"
TODAY = date.today().isoformat()

DOC_SKILL_ROOT = Path(
    r"C:\Users\73921\.codex\plugins\cache\openai-primary-runtime\documents\26.715.12143\skills\documents"
)
PRIVACY_SCRUB = DOC_SKILL_ROOT / "scripts" / "privacy_scrub.py"

INK = RGBColor(23, 50, 77)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 104, 120)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
TEAL = RGBColor(15, 118, 110)
GOLD = RGBColor(151, 111, 20)

FORBIDDEN_VISIBLE = [
    "音翼",
    "音曼",
    "翼欧",
    "张灏",
    "周大",
    "Yinyi",
    "Yinman",
    "Yiou",
    "zhanghao",
    "yinkman",
]

SOURCE_EXCLUDED = {
    "src/data/products.ts",
    "src/features/classroom/brand.ts",
    "src/features/classroom/data/initialProfile.ts",
    "src/lib/aiAdvisor.ts",
}


def rgb(hex_value: str) -> RGBColor:
    hex_value = hex_value.lstrip("#")
    return RGBColor.from_string(hex_value.upper())


def set_run_font(
    run,
    *,
    name: str = "SimSun",
    size: float | None = None,
    color: RGBColor | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, *, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, margin_value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin_value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: Sequence[int], indent_dxa: int = 120):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_field(paragraph, *, total: int | None = None):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=8.5, color=MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    display = OxmlElement("w:t")
    display.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, display, fld_end])
    suffix = f" 页 / 共 {total} 页" if total is not None else " 页"
    suffix_run = paragraph.add_run(suffix)
    set_run_font(suffix_run, size=8.5, color=MUTED)


def set_document_metadata(doc: Document, title: str, subject: str):
    props = doc.core_properties
    props.title = title
    props.subject = subject
    props.author = ""
    props.last_modified_by = ""
    props.company = "" if hasattr(props, "company") else None
    props.comments = ""
    props.keywords = ""
    props.category = "软件著作权登记材料"


def configure_a4_document(doc: Document, *, cover: bool = False):
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(18)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)
    section.header_distance = Mm(10)
    section.footer_distance = Mm(9)
    section.different_first_page_header_footer = cover

    normal = doc.styles["Normal"]
    normal.font.name = "SimSun"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "SimSun")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "SimSun")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[style_name]
        style.font.name = "SimSun"
        style._element.rPr.rFonts.set(qn("w:ascii"), "SimSun")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "SimSun")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_style in ("List Bullet", "List Number"):
        style = doc.styles[list_style]
        style.font.name = "SimSun"
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_para.paragraph_format.space_after = Pt(0)
    run = header_para.add_run(f"{FULL_NAME} / 软件说明书")
    set_run_font(run, size=8.5, color=MUTED)

    footer = section.footer
    footer_para = footer.paragraphs[0]
    add_page_field(footer_para)


def add_title_cover(doc: Document, subtitle: str):
    for _ in range(5):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    run = kicker.add_run("软件著作权登记材料")
    set_run_font(run, size=11, color=GOLD, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run(SOFTWARE_NAME)
    set_run_font(run, size=27, color=INK, bold=True)

    version = doc.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version.paragraph_format.space_after = Pt(18)
    run = version.add_run(VERSION)
    set_run_font(run, size=18, color=BLUE, bold=True)

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_para.paragraph_format.space_after = Pt(72)
    run = subtitle_para.add_run(subtitle)
    set_run_font(run, size=15, color=DARK_BLUE)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.paragraph_format.space_after = Pt(4)
    run = date_para.add_run(f"整理日期：{TODAY}")
    set_run_font(run, size=10.5, color=MUTED)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run("本文件不含申请人以外的人名、单位名或标识")
    set_run_font(run, size=9.5, color=MUTED)
    doc.add_page_break()


def add_body(doc: Document, text: str, *, bold_lead: str | None = None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    if bold_lead and text.startswith(bold_lead):
        first = p.add_run(bold_lead)
        set_run_font(first, size=10.5, bold=True, color=INK)
        rest = p.add_run(text[len(bold_lead) :])
        set_run_font(rest, size=10.5)
    else:
        run = p.add_run(text)
        set_run_font(run, size=10.5)
    return p


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_run_font(run, size=10.5)
    return p


def add_number(doc: Document, text: str):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text)
    set_run_font(run, size=10.5)
    return p


def add_callout(doc: Document, label: str, text: str, *, fill: str = LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    lead = p.add_run(f"{label}  ")
    set_run_font(lead, size=10.5, color=DARK_BLUE, bold=True)
    body = p.add_run(text)
    set_run_font(body, size=10.5)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_figure(doc: Document, image_path: Path, caption: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(6.45))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(2)
    cap.paragraph_format.space_after = Pt(8)
    cap_run = cap.add_run(caption)
    set_run_font(cap_run, size=9, color=MUTED)


def add_key_value_table(doc: Document, rows: Sequence[tuple[str, str]], widths=(2700, 6660)):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    for index, (label, value) in enumerate(rows):
        row = table.rows[0] if index == 0 else table.add_row()
        row.cells[0].text = ""
        row.cells[1].text = ""
        label_run = row.cells[0].paragraphs[0].add_run(label)
        set_run_font(label_run, size=10, bold=True, color=DARK_BLUE)
        value_run = row.cells[1].paragraphs[0].add_run(value)
        set_run_font(value_run, size=10)
        set_cell_shading(row.cells[0], LIGHT_BLUE)
    set_table_geometry(table, list(widths))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def find_font() -> Path | None:
    candidates = [
        Path(r"C:\Windows\Fonts\msyh.ttc"),
        Path(r"C:\Windows\Fonts\simhei.ttf"),
        Path(r"C:\Windows\Fonts\simsun.ttc"),
    ]
    return next((path for path in candidates if path.exists()), None)


def process_screenshots():
    SCREENSHOT_DIR.mkdir(parents=True, exist_ok=True)
    specs = [
        ("presales-form-page.png", "01-presales-parameters.png"),
        ("project-archive-page.png", "02-project-archive.png"),
        ("equipment-list-page.png", "03-equipment-recommendation.png"),
        ("point-map-page.png", "04-point-layout.png"),
        ("topology-page.png", "05-system-topology.png"),
        ("wiring-page.png", "06-interface-wiring.png"),
        ("port-usage-page.png", "07-port-usage.png"),
    ]
    missing = [src for src, _ in specs if not (SCREENSHOT_DIR / src).exists()]
    if missing:
        raise FileNotFoundError(f"Missing screenshot inputs: {', '.join(missing)}")

    font_path = find_font()
    replacement_font = ImageFont.truetype(str(font_path), 16) if font_path else ImageFont.load_default()
    replacement_label_font = ImageFont.truetype(str(font_path), 13) if font_path else ImageFont.load_default()

    for source_name, target_name in specs:
        source = SCREENSHOT_DIR / source_name
        target = SCREENSHOT_DIR / target_name
        with Image.open(source) as image:
            clean = image.convert("RGB").crop((20, 0, min(1038, image.width), min(710, image.height)))
            if source_name == "project-archive-page.png":
                draw = ImageDraw.Draw(clean)
                draw.rectangle((27, 116, 320, 174), fill=(255, 255, 255))
                draw.text((31, 123), "项目", font=replacement_label_font, fill=(86, 100, 95))
                draw.text((31, 146), "示例教室项目", font=replacement_font, fill=(24, 54, 46))
            clean.save(target, "PNG", optimize=True)


def neutralize_source(text: str) -> str:
    replacements = [
        ("音翼科技", "系统品牌甲"),
        ("音翼", "品牌甲"),
        ("音曼", "品牌乙"),
        ("翼欧", "品牌甲"),
        ("张灏", "技术支持"),
        ("周大", "项目联系人"),
        ("Microsoft YaHei", "SimSun"),
        ("Microsoft", "SystemVendor"),
        ("OpenAI", "AIService"),
        ("GitHub", "CodeRepository"),
        ("Google", "WebPlatform"),
        ("©", ""),
    ]
    for old, new in replacements:
        text = text.replace(old, new)

    text = re.sub(r"Yinman", "BrandB", text, flags=re.IGNORECASE)
    text = re.sub(r"Yinyi", "BrandA", text, flags=re.IGNORECASE)
    text = re.sub(r"Yiou", "BrandA", text, flags=re.IGNORECASE)
    text = re.sub(r"zhanghao\w*", "technicalSupport", text, flags=re.IGNORECASE)
    text = re.sub(r"yinkman", "example-service", text, flags=re.IGNORECASE)
    text = re.sub(r"logo", "brandMark", text, flags=re.IGNORECASE)
    text = re.sub(r"copyright", "rightsNotice", text, flags=re.IGNORECASE)
    text = re.sub(r"https?://[^\s\"'`)>]+", "https://service.example.invalid", text)
    text = re.sub(r"[A-Za-z]:\\Users\\[^\\\s]+", r"<local-user>", text)
    text = re.sub(r"/Users/[^/\s]+", r"<local-user>", text)
    return text


def visual_columns(text: str) -> int:
    return sum(2 if ord(char) > 127 else 1 for char in text)


def split_at_visual_width(text: str, max_columns: int) -> tuple[str, str]:
    columns = 0
    hard_index = len(text)
    for index, char in enumerate(text):
        columns += 2 if ord(char) > 127 else 1
        if columns > max_columns:
            hard_index = index
            break
    if hard_index == len(text):
        return text, ""
    floor = max(1, int(hard_index * 0.62))
    preferred = " ,.;:(){}[]<>+-=*/?&|"
    split_index = -1
    for index in range(hard_index, floor, -1):
        if text[index - 1] in preferred:
            split_index = index
            break
    if split_index < 1:
        split_index = hard_index
    return text[:split_index].rstrip(), text[split_index:].lstrip()


def wrap_code_line(line: str, max_columns: int = 166) -> list[str]:
    if visual_columns(line) <= max_columns:
        return [line]
    indent_match = re.match(r"\s*", line)
    base_indent = indent_match.group(0) if indent_match else ""
    continuation_indent = base_indent + "  "
    remaining = line
    result: list[str] = []
    while visual_columns(remaining) > max_columns:
        head, tail = split_at_visual_width(remaining, max_columns)
        result.append(head)
        remaining = continuation_indent + tail
    result.append(remaining)
    return result


def source_paths() -> list[Path]:
    paths = []
    for path in (ROOT / "src").rglob("*"):
        if not path.is_file() or path.suffix.lower() not in {".ts", ".tsx", ".css"}:
            continue
        relative = path.relative_to(ROOT).as_posix()
        if relative in SOURCE_EXCLUDED:
            continue
        if ".test." in path.name or ".spec." in path.name:
            continue
        paths.append(path)

    priority = {
        "src/types.ts": 0,
        "src/features/classroom/types.ts": 1,
        "src/data/questionFlow.ts": 2,
    }
    return sorted(paths, key=lambda path: (priority.get(path.relative_to(ROOT).as_posix(), 50), path.relative_to(ROOT).as_posix()))


def build_source_corpus() -> tuple[list[str], list[str]]:
    corpus: list[str] = []
    included: list[str] = []
    for path in source_paths():
        relative = path.relative_to(ROOT).as_posix()
        included.append(relative)
        corpus.extend(
            [
                "/**",
                f" * File: {neutralize_source(relative)}",
                " */",
            ]
        )
        source = path.read_text(encoding="utf-8-sig")
        for raw_line in source.splitlines():
            clean = neutralize_source(raw_line.rstrip())
            corpus.extend(wrap_code_line(clean))
        corpus.append("")

    if len(corpus) < 6000:
        raise RuntimeError(f"Sanitized source corpus has only {len(corpus)} display lines; 6000 required")
    for forbidden in FORBIDDEN_VISIBLE:
        hits = [index for index, line in enumerate(corpus, 1) if forbidden.lower() in line.lower()]
        if hits:
            raise RuntimeError(f"Forbidden source marker {forbidden!r} remains at lines {hits[:5]}")
    return corpus, included


def configure_code_document(doc: Document, volume_label: str):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Mm(297)
    section.page_height = Mm(210)
    section.top_margin = Mm(10)
    section.bottom_margin = Mm(10)
    section.left_margin = Mm(11)
    section.right_margin = Mm(11)
    section.header_distance = Mm(5)
    section.footer_distance = Mm(5)

    normal = doc.styles["Normal"]
    normal.font.name = "Courier New"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Courier New")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Courier New")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    normal.font.size = Pt(6.4)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    header = section.header
    p = header.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"{FULL_NAME}  源程序代码（{volume_label}）")
    set_run_font(run, name="SimSun", size=8.5, color=MUTED, bold=True)

    footer = section.footer
    footer_p = footer.paragraphs[0]
    add_page_field(footer_p, total=60)


def create_code_document(
    path: Path,
    lines: Sequence[tuple[int, str]],
    *,
    volume_label: str,
):
    if len(lines) != 3000:
        raise ValueError(f"Expected 3000 lines for {volume_label}, got {len(lines)}")
    doc = Document()
    configure_code_document(doc, volume_label)
    set_document_metadata(doc, f"{FULL_NAME} 源程序代码（{volume_label}）", "软件著作权源程序代码文档")

    for page_index in range(60):
        page_lines = lines[page_index * 50 : (page_index + 1) * 50]
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(8.7)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        for line_index, (source_number, source_text) in enumerate(page_lines):
            run = p.add_run(f"{source_number:06d} | {source_text}")
            set_run_font(run, name="Courier New", size=6.4, color=RGBColor(0, 0, 0))
            run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "SimSun")
            if line_index < 49:
                run.add_break(WD_BREAK.LINE)
        if page_index < 59:
            page_break = doc.add_paragraph()
            page_break.paragraph_format.space_before = Pt(0)
            page_break.paragraph_format.space_after = Pt(0)
            page_break.add_run().add_break(WD_BREAK.PAGE)

    doc.save(path)


def create_manual(path: Path):
    doc = Document()
    configure_a4_document(doc, cover=True)
    set_document_metadata(doc, f"{FULL_NAME} 软件说明书", "软件著作权登记软件说明书")
    add_title_cover(doc, "软件说明书")

    doc.add_heading("文档说明", level=1)
    add_body(doc, f"本说明书用于说明{FULL_NAME}的用途、运行环境、主要功能、操作流程和输出结果。软件面向智能音频项目的售前工程设计工作，将参数采集、设备推荐、图纸生成和报告导出集中在同一工作界面。")
    add_callout(doc, "适用对象", "售前工程师、方案设计人员、项目支持人员以及需要快速形成教室、会议室等空间音频方案的使用者。", fill=LIGHT_BLUE)

    doc.add_heading("目录", level=1)
    toc_items = [
        "1  软件概述",
        "2  运行环境",
        "3  总体操作流程",
        "4  售前参数采集",
        "5  项目档案与风险复核",
        "6  自动设备推荐与手动调整",
        "7  点位图生成与标注",
        "8  系统拓扑图",
        "9  接口接线图与接口占用表",
        "10 报告导入、导出与数据保存",
        "11 校核提示与异常处理",
        "12 典型操作示例",
        "13 软件特点与应用边界",
    ]
    for item in toc_items:
        add_bullet(doc, item)
    doc.add_page_break()

    doc.add_heading("1 软件概述", level=1)
    doc.add_heading("1.1 建设目的", level=2)
    add_body(doc, "传统售前方案往往需要在表格、绘图工具和文档之间反复切换。软件通过结构化采集项目参数，自动生成设备清单、阵列麦与音箱点位图、系统拓扑图、接口接线图及接口占用表，减少重复录入并提高方案表达的一致性。")
    doc.add_heading("1.2 主要功能", level=2)
    for text in [
        "采集使用场景、使用需求、房间尺寸、现场条件、声学因素和外接设备信息。",
        "根据项目参数形成设备推荐结果，并允许在受控范围内调整设备数量或方案类型。",
        "输出阵列麦与音箱点位图，显示设备位置、间距、安装高度和方向等工程标注。",
        "输出系统拓扑图，展示设备之间的信号方向、线材类型和数量。",
        "输出接口接线图与接口占用表，帮助核对线材、端口和连接方式。",
        "导入、导出项目数据和完整方案报告，便于归档、交接和复核。",
    ]:
        add_bullet(doc, text)
    doc.add_heading("1.3 软件组成", level=2)
    add_key_value_table(
        doc,
        [
            ("售前采集区", "录入项目基础信息、场景需求、房间参数和现场条件。"),
            ("项目档案区", "汇总关键参数、完整度、声学判断和复勘提醒。"),
            ("方案输出区", "显示设备清单、点位图、拓扑图、接线图和接口占用表。"),
            ("报告功能", "导入既有项目数据，导出包含图纸与清单的方案报告。"),
        ],
    )
    doc.add_page_break()

    doc.add_heading("2 运行环境", level=1)
    add_key_value_table(
        doc,
        [
            ("软件形态", "浏览器运行的图形化应用软件。"),
            ("处理器", "双核及以上处理器。"),
            ("内存", "4 GB 及以上。"),
            ("显示", "建议 1366 x 768 或更高分辨率。"),
            ("浏览器", "支持现代网页标准的桌面端或移动端浏览器。"),
            ("存储", "建议预留 100 MB 以上可用空间，用于页面文件、报告和项目数据。"),
            ("网络", "本地功能可离线使用；如接入在线辅助服务，则需要可用网络。"),
        ],
    )
    doc.add_heading("2.1 启动方式", level=2)
    add_number(doc, "打开软件入口文件或访问部署后的软件地址。")
    add_number(doc, "等待主界面加载完成，确认售前采集、项目档案和方案输出三个区域可见。")
    add_number(doc, "首次使用时从售前采集区开始录入项目参数。")
    add_callout(doc, "提示", "软件会根据参数变化重新计算推荐结果。修改尺寸、需求或现场条件后，应重新核对设备清单和全部图纸。")
    doc.add_page_break()

    doc.add_heading("3 总体操作流程", level=1)
    workflow = [
        ("1. 建立项目", "填写项目名称和客户标识，选择使用场景。"),
        ("2. 采集需求", "选择互动、扩声、录播等需求，并录入房间长、宽、高。"),
        ("3. 补充条件", "选择吊顶、讲台、空调、材质、吸声和外接设备等信息。"),
        ("4. 复核档案", "查看参数完整度、声学判断和需要复勘的项目。"),
        ("5. 确认设备", "采用自动推荐，或按现场约束进行受控调整。"),
        ("6. 核对图纸", "依次核对点位图、拓扑图、接口接线图和接口占用表。"),
        ("7. 导出归档", "导出报告或项目数据，保存本次方案。"),
    ]
    add_key_value_table(doc, workflow, widths=(2500, 6860))
    doc.add_heading("3.1 数据联动", level=2)
    add_body(doc, "软件的输入与输出保持联动。房间尺寸会影响点位分布；使用需求会影响设备类别；已有设备会影响接口和线材；设备数量变化会同步更新清单、拓扑、接线和报告。")
    doc.add_page_break()

    doc.add_heading("4 售前参数采集", level=1)
    add_body(doc, "售前采集区按业务顺序组织参数。带选中状态的按钮表示当前方案条件；数值输入用于房间尺寸和可选的实测声学数据；下拉项用于安装条件和材质信息。")
    add_figure(doc, SCREENSHOT_DIR / "01-presales-parameters.png", "图 1  售前参数采集界面（匿名化示例）")
    doc.add_heading("4.1 场景与需求", level=2)
    add_body(doc, "先选择会议室、普通教室、阶梯教室、报告厅、合班教室或其他场景，再选择互动课堂、本地扩声、录播等需求。部分需求可以组合选择。")
    doc.add_heading("4.2 房间尺寸", level=2)
    add_body(doc, "输入房间长、宽、高。软件据此计算面积和体积，并作为设备数量、覆盖位置、间距和图纸比例的基础参数。")
    doc.add_heading("4.3 现场条件", level=2)
    add_body(doc, "选择吊顶、讲台、空调、地面、顶面、墙面、软装、玻璃比例、桌椅布置和拍手测试等信息。相关参数用于形成声学提示、安装提醒和复勘项。")
    doc.add_page_break()

    doc.add_heading("5 项目档案与风险复核", level=1)
    add_body(doc, "项目档案区将零散输入整理为可复核的项目摘要，显示项目完整度、场景、需求、尺寸、面积、体积、扩声范围、扩声形态以及声学判断。")
    add_figure(doc, SCREENSHOT_DIR / "02-project-archive.png", "图 2  项目档案与完整度检查（匿名化示例）")
    doc.add_heading("5.1 完整度", level=2)
    add_body(doc, "完整度用于提示仍需补充的输入项。未确认的安装条件或复勘项目会以醒目状态显示。")
    doc.add_heading("5.2 复勘提醒", level=2)
    add_body(doc, "复勘提醒只描述参数对结果的影响，例如混响、玻璃、吊顶、已有设备和后场距离对拾音、安装、接口及覆盖的影响。")
    doc.add_page_break()

    doc.add_heading("6 自动设备推荐与手动调整", level=1)
    add_body(doc, "系统根据当前输入形成设备清单，并标记推荐的麦克风和音箱方案。用户可查看数量，也可使用增加、减少和恢复自动推荐功能。")
    add_figure(doc, SCREENSHOT_DIR / "03-equipment-recommendation.png", "图 3  设备推荐、数量调整与恢复推荐")
    doc.add_heading("6.1 自动推荐", level=2)
    add_body(doc, "自动推荐会综合场景、需求、空间尺寸、安装条件和已有设备，形成当前可用的设备组合。")
    doc.add_heading("6.2 手动调整", level=2)
    add_body(doc, "当现场已有明确数量要求时，可通过数量控制按钮调整。手动调整后，应查看点位校核和接口校核结果，避免出现覆盖不足或接口不足。")
    doc.add_heading("6.3 恢复自动推荐", level=2)
    add_body(doc, "点击“恢复自动推荐”后，软件清除设备数量和处理能力的手动覆盖，重新按照当前项目参数计算设备清单。")
    doc.add_page_break()

    doc.add_heading("7 点位图生成与标注", level=1)
    add_body(doc, "点位图用于展示房间边界、讲台位置、阵列麦、音箱和覆盖区域。图中同时给出点位编号、设备间距、距墙距离、安装高度和方向等标注。")
    add_figure(doc, SCREENSHOT_DIR / "04-point-layout.png", "图 4  阵列麦与音箱点位图")
    doc.add_heading("7.1 图纸内容", level=2)
    for text in [
        "房间尺寸和讲台方向。",
        "阵列麦编号、位置和相邻间距。",
        "音箱组编号、安装高度、水平角度和俯角。",
        "阵列麦与音箱的示意覆盖范围。",
        "已有音箱标注和点位增减工具。",
    ]:
        add_bullet(doc, text)
    doc.add_heading("7.2 图纸导出", level=2)
    add_body(doc, "点击点位图区域的导出按钮，可保存当前图纸。导出前应确认所有点位、编号和标注均与设备清单一致。")
    doc.add_page_break()

    doc.add_heading("8 系统拓扑图", level=1)
    add_body(doc, "系统拓扑图从设备关系角度展示信号流向。节点代表设备，带方向的连线表示连接关系，连线文字说明线材类型及数量。")
    add_figure(doc, SCREENSHOT_DIR / "05-system-topology.png", "图 5  系统拓扑图")
    doc.add_heading("8.1 阅读方法", level=2)
    add_number(doc, "先识别阵列麦、处理设备、功放、音箱和外接设备。")
    add_number(doc, "沿箭头方向查看信号从输入设备到输出设备的传递路径。")
    add_number(doc, "核对连线文字中的线材类型和数量。")
    add_number(doc, "将拓扑中的设备数量与设备清单、点位图进行交叉核对。")
    doc.add_heading("8.2 关系同步", level=2)
    add_body(doc, "设备清单或连接关系变化后，拓扑图会同步更新。拓扑图中的线材名称和颜色与接口接线图保持一致。")
    doc.add_page_break()

    doc.add_heading("9 接口接线图与接口占用表", level=1)
    doc.add_heading("9.1 接口接线图", level=2)
    add_body(doc, "接口接线图以设备背面接口为基础，显示每根线的起点、终点和图中编号。鼠标指向线材或端口时，可突出对应线材与两端接口，便于追踪连接关系。")
    add_figure(doc, SCREENSHOT_DIR / "06-interface-wiring.png", "图 6  接口接线图")
    add_body(doc, "对于凤凰端子，图纸显示端子上的实际接线位置；对于带成品接头的线材，图纸显示接头形式。线芯分叉只在需要现场接线的端子侧显示。")
    doc.add_page_break()

    doc.add_heading("9.2 接口占用表", level=2)
    add_body(doc, "接口占用表与接线图编号一一对应。每根线占一行，列出设备、接口、接口形式、线材和接线方式。")
    add_figure(doc, SCREENSHOT_DIR / "07-port-usage.png", "图 7  接口占用表")
    doc.add_heading("9.3 接线核对步骤", level=2)
    add_number(doc, "按图中编号在接口占用表中找到对应记录。")
    add_number(doc, "核对起点设备、终点设备以及两端接口名称。")
    add_number(doc, "核对线材类型、接头形式和端子极性。")
    add_number(doc, "完成一根线后再核对下一编号，避免遗漏或串线。")
    doc.add_page_break()

    doc.add_heading("10 报告导入、导出与数据保存", level=1)
    doc.add_heading("10.1 导出报告", level=2)
    add_body(doc, "点击主界面的“导出报告”按钮，软件汇总项目档案、设备清单、点位图、系统拓扑图、接口接线图、接口占用表和复勘提醒，生成可归档的方案报告。")
    add_bullet(doc, "导出前确认项目名称、客户标识和房间尺寸。")
    add_bullet(doc, "检查设备清单数量与图纸节点数量。")
    add_bullet(doc, "检查接口接线图和接口占用表编号一致。")
    add_bullet(doc, "报告中的图纸采用适合放大查看的清晰度输出。")
    doc.add_heading("10.2 导入报告", level=2)
    add_body(doc, "点击“导入报告”选择由本软件导出的项目数据，软件读取参数并恢复项目状态。导入后应核对页面中的项目名称、尺寸、需求和设备数量。")
    doc.add_heading("10.3 本地草稿", level=2)
    add_body(doc, "开发工作页会将售前采集参数和手动数量调整保存在本地浏览器存储中，刷新页面后可继续工作。对外发布版本首次打开时使用干净的采集表，不读取开发草稿。")
    doc.add_page_break()

    doc.add_heading("11 校核提示与异常处理", level=1)
    add_key_value_table(
        doc,
        [
            ("输入不完整", "返回售前采集区，补充带提醒的必填项。"),
            ("点位待复核", "检查房间尺寸、覆盖目标、安装条件和手动数量。"),
            ("接口不足", "检查设备数量、外接设备和处理设备选择。"),
            ("报告未生成", "确认浏览器允许下载，并重新执行导出。"),
            ("导入失败", "确认文件来自本软件且文件内容完整。"),
            ("页面显示异常", "刷新页面；如问题仍存在，记录项目参数和异常位置。"),
        ],
    )
    doc.add_heading("11.1 校核原则", level=2)
    add_body(doc, "校核结果用于提醒用户检查工程条件，不替代现场勘察。涉及已有系统、复杂接口、特殊安装结构或不确定的设备能力时，应由专业技术人员复核。")
    doc.add_heading("11.2 数据一致性", level=2)
    add_body(doc, "设备清单、点位图、拓扑图、接线图、接口占用表和导出报告应基于同一组项目参数。任何手动调整都需要在全部输出中重新核对。")
    doc.add_page_break()

    doc.add_heading("12 典型操作示例", level=1)
    example_steps = [
        "新建一个普通教室项目，选择互动课堂和本地扩声需求。",
        "输入房间长、宽、高，选择吊顶、讲台、空调和声学条件。",
        "按现场情况选择录播主机、讲台电脑、无线手持麦等外接设备。",
        "查看项目档案完整度，补齐未确认的安装条件。",
        "采用系统推荐的麦克风和音箱方案，核对设备清单。",
        "检查点位图中的编号、间距、安装高度和方向。",
        "检查拓扑图中的信号方向、线材类型和数量。",
        "按编号检查接口接线图和接口占用表。",
        "导出报告并保存项目数据。",
    ]
    for step in example_steps:
        add_number(doc, step)
    add_callout(doc, "完成标准", "项目参数完整；设备清单与各图纸数量一致；接口编号可追踪；报告可正常打开并放大查看。", fill=LIGHT_BLUE)
    doc.add_page_break()

    doc.add_heading("13 软件特点与应用边界", level=1)
    doc.add_heading("13.1 软件特点", level=2)
    for text in [
        "参数、设备、图纸和报告在同一页面联动。",
        "采用结构化规则生成工程输出，便于复核和持续校准。",
        "点位、拓扑、接线和接口占用表相互对应。",
        "支持手动数量调整、恢复自动推荐和项目数据导入导出。",
        "支持桌面宽屏和移动窄屏显示。",
    ]:
        add_bullet(doc, text)
    doc.add_heading("13.2 应用边界", level=2)
    add_body(doc, "软件用于售前阶段的方案辅助设计。最终设备选型、安装位置、供电、负载、接口适配、施工安全和现场调试仍需结合实物资料、现场勘察和专业复核。")
    doc.add_heading("13.3 版本说明", level=2)
    add_key_value_table(
        doc,
        [
            ("软件名称", SOFTWARE_NAME),
            ("软件简称", SHORT_NAME),
            ("版本号", VERSION),
            ("文档用途", "软件著作权登记软件说明书"),
        ],
    )
    add_callout(doc, "文档结束", "以上为本软件的主要功能、操作方法和输出说明。")

    doc.save(path)


def create_basic_info(path: Path):
    doc = Document()
    configure_a4_document(doc, cover=True)
    set_document_metadata(doc, f"{FULL_NAME} 软著基本信息与提交清单", "软件著作权登记准备清单")
    add_title_cover(doc, "基本信息与提交清单")

    doc.add_heading("1 推荐登记名称", level=1)
    add_key_value_table(
        doc,
        [
            ("软件全称", SOFTWARE_NAME),
            ("软件简称", SHORT_NAME),
            ("版本号", VERSION),
            ("作品类型", "应用软件"),
            ("开发方式", "独立开发（提交前由申请人确认）"),
            ("权利取得方式", "原始取得（提交前由申请人确认）"),
        ],
    )
    add_callout(doc, "名称口径", "全部说明书、源代码文档和登记申请表必须使用同一软件名称与版本号。提交前如需改名，应同步修改整套材料。", fill=LIGHT_BLUE)

    doc.add_heading("2 软件功能摘要", level=1)
    add_body(doc, f"{FULL_NAME}是一套面向智能音频项目售前工程设计的应用软件。软件采集使用场景、需求、房间尺寸、现场条件、声学因素和已有设备，自动形成设备清单、阵列麦与音箱点位图、系统拓扑图、接口接线图、接口占用表及方案报告，并支持项目数据导入导出和受控的设备数量调整。")

    doc.add_heading("3 建议技术信息", level=1)
    add_key_value_table(
        doc,
        [
            ("开发语言", "TypeScript、TSX、CSS"),
            ("软件架构", "浏览器端单页应用"),
            ("主要功能模块", "参数采集、推荐计算、图纸生成、接口校核、报告导出、数据保存"),
            ("运行方式", "桌面端或移动端现代浏览器"),
            ("代码规模", "安全匿名化源码超过 6000 展示行，本次提交前、后各 3000 行"),
        ],
    )
    doc.add_page_break()

    doc.add_heading("4 申请人需补充的信息", level=1)
    rows = [
        ("申请人名称", "待申请人填写，必须与证件一致。"),
        ("证件类型及号码", "企业填写统一社会信用代码；个人填写有效身份证件信息。"),
        ("联系地址", "待申请人填写。"),
        ("联系人及联系方式", "待申请人填写。"),
        ("开发完成日期", "待申请人确认，应与研发记录相符。"),
        ("首次发表情况", "选择未发表或填写首次发表日期、地点。"),
        ("著作权归属", "如存在职务开发、委托开发或合作开发，应补充相应证明。"),
    ]
    add_key_value_table(doc, rows, widths=(3000, 6360))
    add_callout(doc, "当前留空原因", "你提供的资料中没有申请人身份和完成日期。为避免写入错误主体，本次成品不虚构姓名、单位或证件信息。", fill="FFF8E8")

    doc.add_heading("5 正式提交材料清单", level=1)
    checklist = [
        "软件著作权登记申请表。",
        "申请人身份证明材料。",
        "软件说明书。",
        "源程序代码文档：前 60 页，每页 50 行。",
        "源程序代码文档：后 60 页，每页 50 行，末页含结束语。",
        "如著作权归属涉及合作、委托、职务开发或权利继受，补充相应证明文件。",
    ]
    for item in checklist:
        add_bullet(doc, item)
    doc.add_page_break()

    doc.add_heading("6 本次已整理的文件", level=1)
    add_key_value_table(
        doc,
        [
            ("01", "软著基本信息与提交清单（DOCX、PDF）"),
            ("02", "软件说明书（DOCX、PDF）"),
            ("03", "源代码文档-前60页（DOCX、PDF）"),
            ("04", "源代码文档-后60页（DOCX、PDF）"),
            ("README", "提交前补充信息和使用说明"),
        ],
    )
    doc.add_heading("7 源代码文档规则", level=1)
    add_bullet(doc, "两册均为 60 页，每页严格 50 行，共 6000 行。")
    add_bullet(doc, "前册取匿名化源码序列最前 3000 行。")
    add_bullet(doc, "后册取匿名化源码序列最后 2999 行，并以“// 源程序代码结束”作为最后一行。")
    add_bullet(doc, "代码页使用连续源码行号、等宽字体和固定行距，便于核查。")
    add_bullet(doc, "不包含构建产物、第三方依赖目录、日志、测试输出或图片资产。")
    doc.add_heading("8 匿名化与版权标识检查", level=1)
    add_body(doc, "软件说明书中的界面图片已裁去品牌头，并对示例项目字段进行匿名化处理；源代码文档对品牌标识、个人姓名、单位名称、网址和本机路径进行了中性替换；文档元数据中的作者和最后修改者已清除。")
    add_callout(doc, "提交前复核", "由申请人补齐身份信息后，再检查申请表、说明书、源码文档的名称与版本号是否完全一致。")
    doc.save(path)


def scrub_docx(path: Path):
    if not PRIVACY_SCRUB.exists():
        raise FileNotFoundError(PRIVACY_SCRUB)
    scrubbed = path.with_name(path.stem + "-scrubbed.docx")
    subprocess.run([sys.executable, str(PRIVACY_SCRUB), str(path), "--out", str(scrubbed)], check=True)
    os.replace(scrubbed, path)


def extract_docx_visible_text(path: Path) -> str:
    chunks: list[str] = []
    with zipfile.ZipFile(path) as archive:
        for name in archive.namelist():
            if not name.endswith(".xml"):
                continue
            if not (name.startswith("word/") or name == "docProps/core.xml"):
                continue
            xml = archive.read(name).decode("utf-8", errors="ignore")
            chunks.extend(re.findall(r"<w:t(?:\s[^>]*)?>(.*?)</w:t>", xml, flags=re.DOTALL))
    return "\n".join(chunks)


def validate_docx(path: Path):
    text = extract_docx_visible_text(path)
    remaining = [marker for marker in FORBIDDEN_VISIBLE if marker.lower() in text.lower()]
    if remaining:
        raise RuntimeError(f"Forbidden visible markers in {path.name}: {remaining}")
    with zipfile.ZipFile(path) as archive:
        names = set(archive.namelist())
        if "docProps/custom.xml" in names:
            raise RuntimeError(f"Custom properties remain in {path.name}")
        core = archive.read("docProps/core.xml").decode("utf-8", errors="ignore")
        if re.search(r"<dc:creator>\s*[^<]+", core) or re.search(r"<cp:lastModifiedBy>\s*[^<]+", core):
            raise RuntimeError(f"Personal core metadata remains in {path.name}")


def write_readme(path: Path, corpus_count: int, included_count: int):
    content = f"""{FULL_NAME} - 软件著作权登记材料

一、已生成材料
1. 01-软著基本信息与提交清单.docx / pdf
2. 02-{SOFTWARE_NAME}{VERSION}-软件说明书.docx / pdf
3. 03-源代码文档-前60页.docx / pdf
4. 04-源代码文档-后60页.docx / pdf

二、源码页规则
- 前册：60页，每页50行，共3000行。
- 后册：60页，每页50行，其中最后一行为“// 源程序代码结束”。
- 匿名化源码序列共 {corpus_count} 行，来自 {included_count} 个项目源码文件。

三、提交前必须补充
- 申请人法定名称和证件信息。
- 联系地址、联系人和联系方式。
- 开发完成日期。
- 首次发表状态；如已发表，填写日期和地点。
- 著作权归属证明（仅在合作、委托、职务开发或权利继受时需要）。

四、提交前最后检查
- 登记申请表、说明书和源码文档的软件名称均为“{SOFTWARE_NAME}”，版本均为“{VERSION}”。
- 不要把 output/software-copyright/qa 目录作为正式材料提交。
- 正式提交使用本目录根部的DOCX或PDF文件。
- 申请人信息补齐后，再做一次人名、单位名和标识检查。
"""
    path.write_text(content, encoding="utf-8")


def create_manifest(
    path: Path,
    corpus_count: int,
    included: Sequence[str],
    front_lines: Sequence[tuple[int, str]],
    back_lines: Sequence[tuple[int, str]],
):
    payload = {
        "softwareName": SOFTWARE_NAME,
        "shortName": SHORT_NAME,
        "version": VERSION,
        "generatedAt": TODAY,
        "sourceCorpusDisplayLines": corpus_count,
        "sourceFiles": list(included),
        "frontDocument": {
            "pages": 60,
            "linesPerPage": 50,
            "firstSourceLine": front_lines[0][0],
            "lastSourceLine": front_lines[-1][0],
        },
        "backDocument": {
            "pages": 60,
            "linesPerPage": 50,
            "firstSourceLine": back_lines[0][0],
            "lastSourceLine": back_lines[-1][0],
            "ending": "// 源程序代码结束",
        },
    }
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    QA_DIR.mkdir(parents=True, exist_ok=True)
    process_screenshots()
    corpus, included = build_source_corpus()

    front_lines = [(index + 1, corpus[index]) for index in range(3000)]
    back_start = len(corpus) - 2999
    back_lines = [(index + 1, corpus[index]) for index in range(back_start, len(corpus))]
    back_lines.append((len(corpus) + 1, "// 源程序代码结束"))

    basic_info = OUT / "01-软著基本信息与提交清单.docx"
    manual = OUT / f"02-{SOFTWARE_NAME}{VERSION}-软件说明书.docx"
    source_front = OUT / "03-源代码文档-前60页.docx"
    source_back = OUT / "04-源代码文档-后60页.docx"

    create_basic_info(basic_info)
    create_manual(manual)
    create_code_document(source_front, front_lines, volume_label="前60页")
    create_code_document(source_back, back_lines, volume_label="后60页")

    for docx_path in (basic_info, manual, source_front, source_back):
        scrub_docx(docx_path)
        validate_docx(docx_path)

    write_readme(OUT / "README-提交前补充信息.txt", len(corpus), len(included))
    create_manifest(QA_DIR / "generation-manifest.json", len(corpus), included, front_lines, back_lines)

    print(
        json.dumps(
            {
                "output": str(OUT),
                "sourceCorpusDisplayLines": len(corpus),
                "sourceFiles": len(included),
                "documents": [basic_info.name, manual.name, source_front.name, source_back.name],
            },
            ensure_ascii=False,
            indent=2,
        )
    )


if __name__ == "__main__":
    main()
