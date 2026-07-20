from __future__ import annotations

import argparse
import html
import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas
from reportlab.platypus import (
    Image,
    KeepTogether,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


SOFTWARE_NAME = "智能音频售前工程方案设计软件 V2.0"
FONT_NORMAL = "SoftRegSong"
FONT_BOLD = "SoftRegHei"


def register_fonts():
    normal_candidates = [
        Path(r"C:\Windows\Fonts\simsun.ttc"),
        Path(r"C:\Windows\Fonts\msyh.ttc"),
    ]
    bold_candidates = [
        Path(r"C:\Windows\Fonts\simhei.ttf"),
        Path(r"C:\Windows\Fonts\msyhbd.ttc"),
        Path(r"C:\Windows\Fonts\msyh.ttc"),
    ]
    normal = next((path for path in normal_candidates if path.exists()), None)
    bold = next((path for path in bold_candidates if path.exists()), None)
    if normal is None or bold is None:
        raise FileNotFoundError("Chinese system fonts were not found")
    pdfmetrics.registerFont(TTFont(FONT_NORMAL, str(normal), subfontIndex=0))
    pdfmetrics.registerFont(TTFont(FONT_BOLD, str(bold), subfontIndex=0))
    pdfmetrics.registerFontFamily(
        "SoftRegCJK",
        normal=FONT_NORMAL,
        bold=FONT_BOLD,
        italic=FONT_NORMAL,
        boldItalic=FONT_BOLD,
    )


def iter_body_blocks(document: Document):
    body = document.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield DocxParagraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield DocxTable(child, document)


def paragraph_has_page_break(paragraph: DocxParagraph) -> bool:
    return bool(paragraph._p.xpath('.//w:br[@w:type="page"]'))


def paragraph_image_blobs(paragraph: DocxParagraph) -> list[bytes]:
    result: list[bytes] = []
    for blip in paragraph._p.xpath(".//a:blip"):
        rel_id = blip.get(qn("r:embed"))
        if not rel_id:
            continue
        related = paragraph.part.related_parts.get(rel_id)
        if related is not None and hasattr(related, "blob"):
            result.append(related.blob)
    return result


def run_markup(paragraph: DocxParagraph) -> str:
    chunks: list[str] = []
    for run in paragraph.runs:
        text = html.escape(run.text).replace("\n", "<br/>")
        if not text:
            continue
        if run.bold:
            text = f"<b>{text}</b>"
        if run.italic:
            text = f"<i>{text}</i>"
        chunks.append(text)
    if chunks:
        return "".join(chunks)
    return html.escape(paragraph.text)


def paragraph_alignment(paragraph: DocxParagraph) -> int:
    if paragraph.alignment is None:
        return TA_LEFT
    name = str(paragraph.alignment)
    if "CENTER" in name:
        return TA_CENTER
    if "JUSTIFY" in name:
        return TA_JUSTIFY
    return TA_LEFT


def paragraph_font_size(paragraph: DocxParagraph, default: float) -> float:
    sizes = [run.font.size.pt for run in paragraph.runs if run.font.size is not None]
    return max(sizes) if sizes else default


def paragraph_is_bold(paragraph: DocxParagraph) -> bool:
    runs = [run for run in paragraph.runs if run.text.strip()]
    return bool(runs) and all(run.bold for run in runs)


def paragraph_style_for(paragraph: DocxParagraph, styles: dict[str, ParagraphStyle]) -> ParagraphStyle:
    style_name = paragraph.style.name if paragraph.style is not None else "Normal"
    has_direct_formatting = (
        paragraph.alignment is not None
        or any(
            run.font.size is not None
            or run.bold is not None
            or run.italic is not None
            or run.font.color.rgb is not None
            for run in paragraph.runs
            if run.text.strip()
        )
    )
    if style_name in styles and (style_name != "Normal" or not has_direct_formatting):
        return styles[style_name]
    size = paragraph_font_size(paragraph, 10.5)
    bold = paragraph_is_bold(paragraph)
    return ParagraphStyle(
        f"direct-{size}-{bold}-{paragraph_alignment(paragraph)}",
        parent=styles["Normal"],
        fontName=FONT_BOLD if bold else FONT_NORMAL,
        fontSize=size,
        leading=max(size * 1.25, size + 2),
        alignment=paragraph_alignment(paragraph),
        textColor=colors.HexColor("#17324D") if size >= 14 else colors.black,
        spaceAfter=6,
    )


def make_styles() -> dict[str, ParagraphStyle]:
    sample = getSampleStyleSheet()
    normal = ParagraphStyle(
        "Normal",
        parent=sample["BodyText"],
        fontName=FONT_NORMAL,
        fontSize=10.5,
        leading=13.2,
        textColor=colors.HexColor("#1E293B"),
        spaceAfter=6,
        alignment=TA_LEFT,
    )
    return {
        "Normal": normal,
        "Heading 1": ParagraphStyle(
            "Heading1",
            parent=normal,
            fontName=FONT_BOLD,
            fontSize=16,
            leading=20,
            textColor=colors.HexColor("#2E74B5"),
            spaceBefore=16,
            spaceAfter=9,
            keepWithNext=True,
        ),
        "Heading 2": ParagraphStyle(
            "Heading2",
            parent=normal,
            fontName=FONT_BOLD,
            fontSize=13,
            leading=17,
            textColor=colors.HexColor("#2E74B5"),
            spaceBefore=12,
            spaceAfter=7,
            keepWithNext=True,
        ),
        "Heading 3": ParagraphStyle(
            "Heading3",
            parent=normal,
            fontName=FONT_BOLD,
            fontSize=12,
            leading=15,
            textColor=colors.HexColor("#1F4D78"),
            spaceBefore=10,
            spaceAfter=5,
            keepWithNext=True,
        ),
        "List Bullet": ParagraphStyle(
            "Bullet",
            parent=normal,
            leftIndent=12 * mm,
            firstLineIndent=-5 * mm,
            bulletIndent=5 * mm,
            spaceAfter=4,
        ),
        "List Number": ParagraphStyle(
            "Number",
            parent=normal,
            leftIndent=12 * mm,
            firstLineIndent=-5 * mm,
            spaceAfter=4,
        ),
        "Caption": ParagraphStyle(
            "Caption",
            parent=normal,
            fontSize=9,
            leading=11,
            textColor=colors.HexColor("#5A6878"),
            alignment=TA_CENTER,
            spaceAfter=8,
        ),
    }


def table_col_widths(docx_table: DocxTable, available_width: float) -> list[float]:
    grid_cols = docx_table._tbl.xpath("./w:tblGrid/w:gridCol")
    widths = []
    for grid_col in grid_cols:
        value = grid_col.get(qn("w:w"))
        widths.append(int(value) if value and value.isdigit() else 0)
    if len(widths) != len(docx_table.columns) or not sum(widths):
        return [available_width / len(docx_table.columns)] * len(docx_table.columns)
    total = sum(widths)
    return [available_width * width / total for width in widths]


def docx_table_to_flowable(docx_table: DocxTable, styles, available_width: float):
    rows = []
    for row in docx_table.rows:
        converted = []
        for cell in row.cells:
            text = "<br/>".join(html.escape(paragraph.text) for paragraph in cell.paragraphs if paragraph.text)
            converted.append(Paragraph(text or " ", styles["Normal"]))
        rows.append(converted)
    widths = table_col_widths(docx_table, available_width)
    table = Table(rows, colWidths=widths, repeatRows=0, hAlign="LEFT")
    commands = [
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#CBD5E1")),
        ("INNERGRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#DCE3EB")),
        ("LEFTPADDING", (0, 0), (-1, -1), 7),
        ("RIGHTPADDING", (0, 0), (-1, -1), 7),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]
    for row_index, row in enumerate(docx_table.rows):
        for col_index, cell in enumerate(row.cells):
            fill = cell._tc.xpath("./w:tcPr/w:shd/@w:fill")
            if fill and fill[0] not in {"auto", "FFFFFF"}:
                commands.append(("BACKGROUND", (col_index, row_index), (col_index, row_index), colors.HexColor(f"#{fill[0]}")))
    table.setStyle(TableStyle(commands))
    return table


def render_standard_docx(input_path: Path, output_path: Path):
    document = Document(input_path)
    styles = make_styles()
    page_width, page_height = A4
    left = right = 20 * mm
    top = 20 * mm
    bottom = 18 * mm
    available_width = page_width - left - right
    story = []
    image_index = 0
    number_index = 0
    temp_dir = output_path.parent / (output_path.stem + "-images")
    temp_dir.mkdir(parents=True, exist_ok=True)

    for block in iter_body_blocks(document):
        if isinstance(block, DocxTable):
            story.append(docx_table_to_flowable(block, styles, available_width))
            story.append(Spacer(1, 6))
            continue

        paragraph = block
        if paragraph_has_page_break(paragraph):
            story.append(PageBreak())
            number_index = 0
            continue

        blobs = paragraph_image_blobs(paragraph)
        if blobs:
            for blob in blobs:
                image_index += 1
                image_path = temp_dir / f"image-{image_index:02d}.png"
                image_path.write_bytes(blob)
                flow_image = Image(str(image_path))
                max_width = available_width
                max_height = 210 * mm
                scale = min(max_width / flow_image.imageWidth, max_height / flow_image.imageHeight, 1.0)
                flow_image.drawWidth = flow_image.imageWidth * scale
                flow_image.drawHeight = flow_image.imageHeight * scale
                flow_image.hAlign = "CENTER"
                story.append(KeepTogether([Spacer(1, 4), flow_image, Spacer(1, 4)]))
            continue

        text = paragraph.text.strip()
        if not text:
            after = paragraph.paragraph_format.space_after
            space = after.pt if after is not None else 5
            story.append(Spacer(1, max(3, min(space, 14))))
            continue

        style_name = paragraph.style.name if paragraph.style is not None else "Normal"
        style = paragraph_style_for(paragraph, styles)
        markup = run_markup(paragraph)
        bullet_text = None
        if style_name == "List Bullet":
            bullet_text = "•"
        elif style_name == "List Number":
            number_index += 1
            bullet_text = f"{number_index}."
        elif style_name not in {"List Number", "List Bullet"}:
            number_index = 0
        story.append(Paragraph(markup, style, bulletText=bullet_text))

    def on_page(pdf_canvas, doc):
        page = pdf_canvas.getPageNumber()
        pdf_canvas.saveState()
        if page > 1:
            pdf_canvas.setFont(FONT_NORMAL, 8.5)
            pdf_canvas.setFillColor(colors.HexColor("#5A6878"))
            pdf_canvas.drawString(left, page_height - 11 * mm, f"{SOFTWARE_NAME} / 软件著作权登记材料")
        pdf_canvas.setFont(FONT_NORMAL, 8.5)
        pdf_canvas.setFillColor(colors.HexColor("#5A6878"))
        pdf_canvas.drawRightString(page_width - right, 9 * mm, f"第 {page} 页")
        pdf_canvas.restoreState()

    pdf = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        leftMargin=left,
        rightMargin=right,
        topMargin=top,
        bottomMargin=bottom,
        title=input_path.stem,
        author="",
        subject="软件著作权登记材料",
    )
    pdf.build(story, onFirstPage=on_page, onLaterPages=on_page)


def extract_code_pages(input_path: Path) -> list[list[str]]:
    document = Document(input_path)
    pages = []
    for paragraph in document.paragraphs:
        text = paragraph.text
        if " | " not in text:
            continue
        lines = text.splitlines()
        if len(lines) != 50:
            raise RuntimeError(f"Code paragraph has {len(lines)} lines instead of 50")
        pages.append(lines)
    if len(pages) != 60:
        raise RuntimeError(f"Code document has {len(pages)} pages instead of 60")
    return pages


def fit_code_font_size(lines: list[str], usable_width: float, start_size: float = 6.4) -> float:
    size = start_size
    while size >= 5.0:
        widest = max(pdfmetrics.stringWidth(line, FONT_NORMAL, size) for line in lines)
        if widest <= usable_width:
            return size
        size -= 0.1
    raise RuntimeError("A source line is too wide for the code page")


def render_code_docx(input_path: Path, output_path: Path):
    pages = extract_code_pages(input_path)
    page_width, page_height = landscape(A4)
    left = 11 * mm
    right = 11 * mm
    usable_width = page_width - left - right
    header_y = page_height - 8 * mm
    first_line_y = page_height - 18 * mm
    line_step = 9.15
    pdf_canvas = canvas.Canvas(
        str(output_path),
        pagesize=landscape(A4),
        pageCompression=1,
        invariant=1,
    )
    pdf_canvas.setTitle(input_path.stem)
    pdf_canvas.setAuthor("")
    pdf_canvas.setSubject("软件著作权源程序代码文档")
    volume = "前60页" if "前60页" in input_path.name else "后60页"

    for page_index, lines in enumerate(pages, 1):
        font_size = fit_code_font_size(lines, usable_width)
        pdf_canvas.setFillColor(colors.HexColor("#5A6878"))
        pdf_canvas.setFont(FONT_BOLD, 8.5)
        pdf_canvas.drawString(left, header_y, f"{SOFTWARE_NAME}  源程序代码（{volume}）")
        pdf_canvas.setFont(FONT_NORMAL, 8.2)
        pdf_canvas.drawRightString(page_width - right, 7 * mm, f"第 {page_index} 页 / 共 60 页")

        pdf_canvas.setFillColor(colors.black)
        pdf_canvas.setFont(FONT_NORMAL, font_size)
        y = first_line_y
        for line in lines:
            pdf_canvas.drawString(left, y, line)
            y -= line_step
        pdf_canvas.showPage()

    pdf_canvas.save()


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("input", type=Path)
    parser.add_argument("output", type=Path)
    parser.add_argument("--code", action="store_true")
    args = parser.parse_args()
    register_fonts()
    args.output.parent.mkdir(parents=True, exist_ok=True)
    if args.code:
        render_code_docx(args.input, args.output)
    else:
        render_standard_docx(args.input, args.output)
    print(args.output)


if __name__ == "__main__":
    main()
